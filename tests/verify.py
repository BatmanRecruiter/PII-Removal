"""End-to-end verification harness — builds fixtures in memory and checks redaction.

Run: .venv/bin/python verify.py
"""
import sys
from io import BytesIO

import fitz
from docx import Document

from app.analyzer import analyze_text
from app.redact_docx import redact_docx
from app.redact_pdf import ScannedPDFError, redact_pdf

PII = {
    "name": "Bruce Wayne",
    "email": "bruce@wayne-tech.com",
    "phone": "(202) 456-1111",
    "url": "linkedin.com/in/brucewayne",
    "address": "1007 Mountain Drive, Austin, TX 78701",
}
# These must SURVIVE (resume content we intentionally keep).
# Real city "Austin" in job history must survive (NER tags it GPE, which we keep).
KEEP = ["2015", "2019", "Austin", "Senior Engineer", "May 2020"]

passed, failed = [], []


def check(label, cond):
    (passed if cond else failed).append(label)
    print(f"  [{'PASS' if cond else 'FAIL'}] {label}")


# --- 1. Analyzer unit check -------------------------------------------------
print("\n== Analyzer ==")
text = f"{PII['name']} | {PII['email']} | {PII['phone']} | {PII['url']} | {PII['address']}"
spans = analyze_text(text)
redacted_text = ""
last = 0
for a, b in spans:
    redacted_text += text[last:a] + "X"
    last = b
redacted_text += text[last:]
for key, val in PII.items():
    check(f"analyzer removes {key}", val not in redacted_text)
# A bare date must survive analysis untouched.
date_spans = analyze_text("In 2015 I joined; graduated May 2019.")
check("analyzer keeps dates (no spans)", date_spans == [])


# --- 2. DOCX ----------------------------------------------------------------
print("\n== DOCX ==")
doc = Document()
# Name split across 3 runs (simulates Word splitting).
p = doc.add_paragraph()
p.add_run("Bru")
r = p.add_run("ce Way")
r.bold = True
p.add_run("ne")
doc.add_paragraph(f"Email: {PII['email']}   Phone: {PII['phone']}")
doc.add_paragraph(f"Portfolio: {PII['url']}")
doc.add_paragraph(f"Address: {PII['address']}")
doc.add_paragraph("Experience: Senior Engineer in Austin (May 2020 - 2019, joined 2015)")
# headshot picture that must be stripped
_pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 32, 32))
_pix.clear_with(128)
doc.add_picture(BytesIO(_pix.tobytes("png")))
# table
tbl = doc.add_table(rows=1, cols=1)
tbl.rows[0].cells[0].text = f"Reference: {PII['name']}, {PII['email']}"
# header & footer
sec = doc.sections[0]
sec.header.paragraphs[0].text = f"{PII['name']} - Resume"
sec.footer.paragraphs[0].text = f"Contact: {PII['phone']}"
doc.core_properties.author = "Bruce Wayne"

buf = BytesIO()
doc.save(buf)
out, _docx_warnings = redact_docx(buf.getvalue())

rd = Document(BytesIO(out))
all_text = "\n".join(par.text for par in rd.paragraphs)
for t in rd.tables:
    for row in t.rows:
        for cell in row.cells:
            all_text += "\n" + cell.text
for s in rd.sections:
    all_text += "\n" + s.header.paragraphs[0].text + "\n" + s.footer.paragraphs[0].text

check("docx body name removed (multi-run)", "Wayne" not in all_text and "Bruce" not in all_text)
check("docx email removed", PII["email"] not in all_text)
check("docx phone removed", "456-1111" not in all_text)
check("docx url removed", PII["url"] not in all_text)
check("docx street address removed", "1007 Mountain Drive" not in all_text)
check("docx table PII removed", PII["email"] not in all_text)
check("docx header name removed", "Bruce Wayne - Resume" not in all_text)
check("docx footer phone removed", "456-1111" not in all_text)
check("docx [REDACTED] token present", "[REDACTED]" in all_text)
check("docx keeps year 2015", "2015" in all_text)
check("docx keeps city Austin", "Austin" in all_text)
check("docx keeps 'Senior Engineer'", "Senior Engineer" in all_text)
check("docx metadata author scrubbed", rd.core_properties.author in ("", None))
check("docx embedded picture removed", not rd.inline_shapes)
# bold formatting on the redacted run should still carry the token
check("docx file is valid (re-openable)", len(all_text) > 0)


# --- 3. PDF -----------------------------------------------------------------
print("\n== PDF ==")
pdf = fitz.open()
page = pdf.new_page()
body = (
    f"{PII['name']}\n{PII['email']}   {PII['phone']}\n{PII['url']}\n"
    f"{PII['address']}\nSenior Engineer, Gotham. Joined 2015, promoted 2019. Based in Austin."
)
page.insert_text((72, 72), body, fontsize=11)
pdf_bytes = BytesIO()
pdf.save(pdf_bytes)
pdf.close()

out_pdf, _pdf_warnings = redact_pdf(pdf_bytes.getvalue())
rp = fitz.open(stream=out_pdf, filetype="pdf")
extracted = "\n".join(pg.get_text("text") for pg in rp)

check("pdf name deleted", "Wayne" not in extracted)
check("pdf email deleted", PII["email"] not in extracted)
check("pdf phone deleted", "456-1111" not in extracted)
check("pdf url deleted", "brucewayne" not in extracted)
check("pdf address deleted", "Mountain Drive" not in extracted)
check("pdf keeps year 2015", "2015" in extracted)
check("pdf keeps city Austin", "Austin" in extracted)
check("pdf metadata scrubbed", not (rp.metadata or {}).get("author"))
rp.close()


# --- 3b. Invisible characters, images, links, vector-text warning -----------
print("\n== PDF edge cases ==")
# Zero-width spaces glued to words (Google Docs PDF exports do this) must not
# hide a name from NER.
zw = fitz.open()
zpage = zw.new_page()
zpage.insert_text((72, 72), "References:\n​Harlan​ ​Smith​, Senior Director\nharlan@example.com")
# a picture that must be blacked out
pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 40, 40))
pix.clear_with(200)
zpage.insert_image(fitz.Rect(400, 72, 460, 132), pixmap=pix)
# a link whose visible text is innocent but whose URI is PII
zpage.insert_link({"kind": fitz.LINK_URI, "from": fitz.Rect(72, 150, 130, 165),
                   "uri": "mailto:harlan@example.com"})
zb = BytesIO(); zw.save(zb); zw.close()
zout, zwarn = redact_pdf(zb.getvalue())
zr = fitz.open(stream=zout, filetype="pdf")
ztext = zr[0].get_text("text")
check("pdf name with zero-width spaces deleted", "Harlan" not in ztext and "Smith" not in ztext)
check("pdf embedded image removed", not zr[0].get_images(full=True))
check("pdf mailto link removed", not zr[0].get_links())
check("plain pdf yields no warnings", zwarn == [])
zr.close()

# --- 3c. OCR pass for text drawn as graphics --------------------------------
print("\n== PDF OCR (text drawn as graphics) ==")
from app import config as cfg  # noqa: E402  (late import: tests tweak knobs)


def make_vector_text_pdf() -> bytes:
    """A page whose PII exists only as pixels: real body text (so the scanned
    check passes), bezier blobs (to trip the vector-text heuristic), and an
    image of rendered PII text standing in for letter outlines."""
    # Render PII text into a small pixmap ("text as pixels", like outlines).
    src = fitz.open(); sp = src.new_page(width=400, height=60)
    sp.insert_text((10, 35), "Contact jane.doe@example.com now", fontsize=18)
    pii_png = sp.get_pixmap(dpi=150).tobytes("png"); src.close()

    vec = fitz.open(); vpage = vec.new_page()
    vpage.insert_text((72, 72), "Normal body text with enough characters to not look scanned.")
    vpage.insert_image(fitz.Rect(72, 300, 472, 360), stream=pii_png)
    shape = vpage.new_shape()
    for i in range(60):  # filled bezier blobs standing in for letter outlines (120 curves)
        x = 72 + (i % 30) * 12
        shape.draw_bezier(fitz.Point(x, 200), fitz.Point(x + 4, 190),
                          fitz.Point(x + 8, 210), fitz.Point(x + 10, 200))
        shape.draw_bezier(fitz.Point(x + 10, 200), fitz.Point(x + 8, 214),
                          fitz.Point(x + 4, 214), fitz.Point(x, 200))
        shape.finish(fill=(0, 0, 0), closePath=True)
    shape.commit()
    vb = BytesIO(); vec.save(vb); vec.close()
    return vb.getvalue()


# With images kept, only OCR can find the PII "drawn" on the page.
_saved_images = cfg.REDACT_IMAGES
cfg.REDACT_IMAGES = False
try:
    vout, vwarn = redact_pdf(make_vector_text_pdf())
    check("ocr ran, no warning", vwarn == [])
    vr = fitz.open(stream=vout, filetype="pdf")
    vp = vr.load_page(0)
    tp = vp.get_textpage_ocr(flags=0, language="eng", dpi=200, full=True,
                             tessdata=cfg.TESSDATA_DIR)
    ocr_out = vp.get_text("text", textpage=tp)
    check("ocr redacted graphics-only PII", "example.com" not in ocr_out
          and "jane.doe" not in ocr_out)
    check("ocr left normal body text alone", "Normal body text" in ocr_out)
    del tp
    vr.close()
finally:
    cfg.REDACT_IMAGES = _saved_images

# When OCR can't run (no tessdata), the same file must produce a warning.
_saved_tess = cfg.TESSDATA_DIR
cfg.TESSDATA_DIR = "/nonexistent-tessdata"
try:
    _wout, wwarn = redact_pdf(make_vector_text_pdf())
    check("missing tessdata falls back to warning",
          any("OCR could not run" in w for w in wwarn))
finally:
    cfg.TESSDATA_DIR = _saved_tess


# --- 4. Scanned PDF rejection ----------------------------------------------
print("\n== Scanned PDF ==")
scan = fitz.open()
scan.new_page()  # blank page, no text layer
scan_bytes = BytesIO()
scan.save(scan_bytes)
scan.close()
try:
    redact_pdf(scan_bytes.getvalue())
    check("scanned pdf rejected", False)
except ScannedPDFError:
    check("scanned pdf rejected", True)


# --- summary ----------------------------------------------------------------
print(f"\n{'='*40}\n{len(passed)} passed, {len(failed)} failed")
if failed:
    print("FAILED:", ", ".join(failed))
    sys.exit(1)
print("ALL GREEN")

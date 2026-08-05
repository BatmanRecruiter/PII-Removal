"""In-process HTTP verification of the full FastAPI app via Starlette TestClient.

Exercises real routing, multipart upload, validation, the async job flow
(submit -> poll -> download), streaming response, and Content-Disposition
headers. Run: .venv/bin/python -m tests.verify_http
"""
import sys
import time
from io import BytesIO

import fitz
from docx import Document
from fastapi.testclient import TestClient

from urllib.parse import unquote
from app.main import app

client = TestClient(app)
passed, failed = [], []


def check(label, cond):
    (passed if cond else failed).append(label)
    print(f"  [{'PASS' if cond else 'FAIL'}] {label}")


def make_docx() -> bytes:
    d = Document()
    p = d.add_paragraph()
    p.add_run("Bru"); p.add_run("ce Way").bold = True; p.add_run("ne")
    d.add_paragraph("Email: bruce@wayne-tech.com  Phone: (202) 456-1111")
    d.add_paragraph("Senior Engineer in Austin, joined 2015")
    b = BytesIO(); d.save(b); return b.getvalue()


def make_pdf() -> bytes:
    doc = fitz.open(); pg = doc.new_page()
    pg.insert_text((72, 72), "Bruce Wayne\nbruce@wayne-tech.com (202) 456-1111\nAustin 2015")
    b = BytesIO(); doc.save(b); doc.close(); return b.getvalue()


def run_job(name: str, content: bytes, mime: str, timeout: float = 60.0):
    """Submit a file and poll to completion.

    Returns (final status response, download response, job_id)."""
    r = client.post("/redact", files={"file": (name, content, mime)})
    if r.status_code != 200:
        return r, None, None
    job_id = r.json()["job_id"]
    deadline = time.time() + timeout
    while time.time() < deadline:
        s = client.get(f"/jobs/{job_id}")
        if s.status_code != 200 or s.json().get("state") in ("done", "error"):
            dl = None
            if s.status_code == 200 and s.json().get("state") == "done":
                dl = client.get(f"/jobs/{job_id}/download")
            return s, dl, job_id
        time.sleep(0.2)
    return s, None, job_id


print("== GET / ==")
r = client.get("/")
check("index returns 200 html", r.status_code == 200 and "PII Redactor" in r.text)
check("healthz ok", client.get("/healthz").json() == {"status": "ok"})

print("\n== POST /redact (docx job) ==")
status, dl, _ = run_job("Resume.docx", make_docx(),
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
check("docx job done", status.status_code == 200 and status.json()["state"] == "done")
check("docx download 200", dl is not None and dl.status_code == 200)
check("docx filename suffixed", "Resume - Redacted.docx" in unquote(dl.headers.get("content-disposition", "")))
check("docx mime correct", "wordprocessingml" in dl.headers.get("content-type", ""))
rd = Document(BytesIO(dl.content))
txt = "\n".join(p.text for p in rd.paragraphs)
check("docx PII gone (name/email/phone)",
      "Wayne" not in txt and "wayne-tech" not in txt and "456-1111" not in txt)
check("docx keeps Austin + 2015", "Austin" in txt and "2015" in txt)
check("docx token present", "[REDACTED]" in txt)

print("\n== POST /redact (pdf job) ==")
status, dl, pdf_job_id = run_job("CV.pdf", make_pdf(), "application/pdf")
check("pdf job done", status.status_code == 200 and status.json()["state"] == "done")
check("pdf no warnings", status.json().get("warnings") == [])
check("pdf download 200", dl is not None and dl.status_code == 200)
check("pdf filename suffixed", "CV - Redacted.pdf" in unquote(dl.headers.get("content-disposition", "")))
rp = fitz.open(stream=dl.content, filetype="pdf")
ext = "\n".join(pg.get_text() for pg in rp); rp.close()
check("pdf PII deleted", "Wayne" not in ext and "wayne-tech" not in ext and "456-1111" not in ext)
check("pdf keeps Austin + 2015", "Austin" in ext and "2015" in ext)
check("download is one-shot (second fetch 404)",
      client.get(f"/jobs/{pdf_job_id}/download").status_code == 404)

print("\n== Error handling ==")
check("rejects .txt (415)",
      client.post("/redact", files={"file": ("x.txt", b"hi", "text/plain")}).status_code == 415)
check("rejects fake docx content (415)",
      client.post("/redact", files={"file": ("x.docx", b"not a zip", "application/octet-stream")}).status_code == 415)
check("rejects empty file (400)",
      client.post("/redact", files={"file": ("x.pdf", b"", "application/pdf")}).status_code == 400)
check("unknown job -> 404", client.get("/jobs/deadbeef").status_code == 404)
# scanned pdf -> job ends in a 422 error status
sd = fitz.open(); sd.new_page(); sb = BytesIO(); sd.save(sb); sd.close()
status, _, _ = run_job("scan.pdf", sb.getvalue(), "application/pdf")
check("scanned pdf -> 422 json error",
      status.status_code == 422 and "error" in status.json())

print(f"\n{'='*40}\n{len(passed)} passed, {len(failed)} failed")
if failed:
    print("FAILED:", ", ".join(failed)); sys.exit(1)
print("ALL GREEN")
